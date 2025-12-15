import os
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from config import config

class SermonFormatter:
    def __init__(self):
        self.config = config
        self.metadata = self.config.get('metadata')
        self.output_configs = self.config.get('output', 'formats')
        
    def format_sermon(self, transcription: str, metadata: Optional[Dict] = None) -> Dict[str, str]:
        """Format the sermon transcription according to the configuration.
        
        Args:
            transcription: Raw transcription text
            metadata: Additional metadata to include in the formatted output
            
        Returns:
            Dict with file paths for each output format
        """
        if metadata is None:
            metadata = {}
            
        # Merge default metadata with provided metadata
        metadata = {**self.metadata, **metadata}
        
        # Generate base filename
        filename = datetime.now().strftime(
            self.config.get('output', 'naming_convention')
        )
        
        # Create formatted content
        formatted_content = self._apply_formatting(transcription, metadata)
        
        # Save in all requested formats
        output_files = {}
        for fmt in self.output_configs:
            output_files[fmt['type']] = self._save_format(
                fmt, filename, formatted_content, metadata
            )
            
        return output_files
    
    def _apply_formatting(self, text: str, metadata: Dict) -> Dict[str, Any]:
        """Apply formatting to the raw transcription."""
        # This is a simple formatter - you can enhance this with more sophisticated logic
        sections = {}
        
        # Get the formatting configuration
        formatting = self.config.get('post_processing', 'formatting')
        
        # Apply text cleanup
        if self.config.get('post_processing', 'text_cleanup', 'capitalize_sentences'):
            text = self._capitalize_sentences(text)
            
        if self.config.get('post_processing', 'text_cleanup', 'fix_punctuation'):
            text = self._fix_punctuation(text)
        
        # For now, just return the text with some basic formatting
        # In a real implementation, you would parse the text and structure it
        # according to the sections in the config
        return {
            'title': f"Sermon - {metadata.get('default_date', datetime.now().strftime('%Y-%m-%d'))}",
            'content': text,
            'metadata': metadata
        }
    
    def _save_format(self, fmt: Dict, filename: str, content: Dict, metadata: Dict) -> str:
        """Save content in the specified format."""
        save_dir = Path(fmt['save_directory'])
        save_dir.mkdir(parents=True, exist_ok=True)
        
        if fmt['type'] == 'txt':
            return self._save_txt(save_dir, filename, content)
        elif fmt['type'] == 'docx':
            return self._save_docx(save_dir, filename, content, fmt.get('template', {}))
        else:
            raise ValueError(f"Unsupported format: {fmt['type']}")
    
    def _save_txt(self, directory: Path, filename: str, content: Dict) -> str:
        """Save content as a text file."""
        filepath = directory / f"{filename}.txt"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Write title
            f.write(f"{content['title']}\n")
            f.write("=" * len(content['title']) + "\n\n")
            
            # Write metadata
            if 'metadata' in content:
                f.write("METADATA\n")
                f.write("-" * 8 + "\n")
                for key, value in content['metadata'].items():
                    if isinstance(value, dict):
                        f.write(f"{key}:\n")
                        for k, v in value.items():
                            f.write(f"  {k}: {v}\n")
                    else:
                        f.write(f"{key}: {value}\n")
                f.write("\n")
            
            # Write content
            f.write("TRANSCRIPTION\n")
            f.write("-" * 12 + "\n")
            f.write(content['content'] + "\n")
            
        return str(filepath)
    
    def _save_docx(self, directory: Path, filename: str, content: Dict, template: Dict) -> str:
        """Save content as a Word document."""
        filepath = directory / f"{filename}.docx"
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = template.get('font_name', 'Times New Roman')
        font.size = Pt(template.get('font_size', 12))
        
        # Add title
        title = doc.add_heading(content['title'], level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata section
        if 'metadata' in content:
            doc.add_heading('Sermon Details', level=2)
            for key, value in content['metadata'].items():
                if isinstance(value, dict):
                    doc.add_paragraph(f"{key}:")
                    for k, v in value.items():
                        doc.add_paragraph(f"  {k}: {v}", style='List Bullet')
                else:
                    doc.add_paragraph(f"{key}: {value}")
            
            doc.add_page_break()
        
        # Add content
        doc.add_heading('Sermon Transcription', level=2)
        
        # Split content into paragraphs and add to document
        paragraphs = content['content'].split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        # Save the document
        doc.save(str(filepath))
        return str(filepath)
    
    def _capitalize_sentences(self, text: str) -> str:
        """Capitalize the first letter of each sentence."""
        if not text:
            return text
            
        sentences = text.split('. ')
        sentences = [s[0].upper() + s[1:] if s else s for s in sentences]
        return '. '.join(sentences)
    
    def _fix_punctuation(self, text: str) -> str:
        """Fix common punctuation issues."""
        if not text:
            return text
            
        # Add space after period if missing
        text = text.replace('.', '. ')
        
        # Remove double spaces
        while '  ' in text:
            text = text.replace('  ', ' ')
            
        return text.strip()
